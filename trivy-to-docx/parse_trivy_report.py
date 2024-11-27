import json
import argparse
import os
from docx import Document

def get_packages_from_trivy_report(json_report):
    # Load the JSON data
    with open(json_report, 'r') as file:
        data = json.load(file)

    packages_info = []

    # Traverse through the report and collect package info
    if 'Results' in data:
        for result in data['Results']:
            if 'Vulnerabilities' in result:
                for vuln in result['Vulnerabilities']:
                    pkg_name = vuln.get('PkgName', 'N/A')
                    installed_version = vuln.get('InstalledVersion', 'N/A')
                    cve_id = vuln.get('VulnerabilityID', 'N/A')
                    severity = vuln.get('Severity', 'Unknown')

                    # Format the CVE with severity
                    cve_with_severity = f"{cve_id} ({severity})"

                    # Find if the package is already in the list
                    found = False
                    for pkg in packages_info:
                        if pkg['name'] == pkg_name and pkg['version'] == installed_version:
                            pkg['vulnerabilities'].append(cve_with_severity)
                            found = True
                            break

                    # If not found, add a new entry
                    if not found:
                        packages_info.append({
                            'name': pkg_name,
                            'version': installed_version,
                            'vulnerabilities': [cve_with_severity]
                        })

    return packages_info

def create_doc_with_table(packages_info, output_docx):
    # Create a new Document
    doc = Document()

    # Add a title to the document
    doc.add_heading('Trivy Report - Vulnerability Summary', level=1)

    # Add a table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Add table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Package Name'
    hdr_cells[1].text = 'Installed Version'
    hdr_cells[2].text = 'Vulnerabilities (CVE with Severity)'

    # Add rows for each package
    for pkg in packages_info:
        row_cells = table.add_row().cells
        row_cells[0].text = pkg['name']
        row_cells[1].text = pkg['version']

        # Adding vulnerabilities with severity in separate lines
        vulnerabilities_text = '\n'.join(pkg['vulnerabilities'])
        row_cells[2].text = vulnerabilities_text

    # Save the document
    doc.save(output_docx)

def main():
    # Set up argument parsing
    parser = argparse.ArgumentParser(description="Generate a DOCX report from a Trivy JSON report.")
    parser.add_argument(
        'trivy_report',
        help="Path to the Trivy JSON report file."
    )

    # Parse arguments
    args = parser.parse_args()

    # Derive the output file name
    input_file = args.trivy_report
    base_name, _ = os.path.splitext(input_file)
    output_file = f"{base_name}.docx"

    # Process the Trivy report and generate the Word document
    packages_info = get_packages_from_trivy_report(args.trivy_report)
    create_doc_with_table(packages_info, output_file)

    print(f"Document '{output_file}' created successfully!")

if __name__ == "__main__":
    main()
